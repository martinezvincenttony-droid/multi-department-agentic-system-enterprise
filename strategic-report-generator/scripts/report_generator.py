#!/usr/bin/env python3
"""Strategic Report Generator - Professional document creation engine."""

import os
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


# ── Brand Colors ────────────────────────────────────────────────
# Default colors for professional reports
PRIMARY_BLUE = RGBColor(0x2F, 0x54, 0x96) # Example blue
ACCENT_GRAY = RGBColor(0x40, 0x40, 0x40) # Example gray
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0) # Example light blue
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x27, 0xAE, 0x60)
AMBER = RGBColor(0xF3, 0x9C, 0x12)


def create_base_document():
    """Create a base document with professional styling."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = ACCENT_GRAY

    # Configure heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = PRIMARY_BLUE
        heading_style.font.name = "Calibri"
        if level == 1:
            heading_style.font.size = Pt(18)
        elif level == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    return doc


def add_title_page(doc, title, subtitle="", author="the organization", date=None):
    """Add a professional title page."""
    if date is None:
        date = datetime.now().strftime("%B %d, %Y")

    # Add spacing before title
    for _ in range(6):
        doc.add_paragraph()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = PRIMARY_BLUE
    run.font.bold = True
    run.font.name = "Calibri"

    # Horizontal line
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_para.add_run("\u2500" * 50)
    run.font.color.rgb = LIGHT_BLUE
    run.font.size = Pt(12)

    # Subtitle
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = ACCENT_GRAY
        run.font.name = "Calibri"

    # Spacing
    doc.add_paragraph()
    doc.add_paragraph()

    # Author and Date
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"Prepared by: {author}")
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT_GRAY

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Date: {date}")
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT_GRAY

    # Page break
    doc.add_page_break()


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = str(header)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        # Set header background
        shading_elm = cell._element.get_or_add_tcPr()
        shading = shading_elm.makeelement(qn("w:shd"), {
            qn("w:fill"): "2F5496",
            qn("w:val"): "clear",
        })
        shading_elm.append(shading)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(value) if value is not None else ""
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run()
            run.font.size = Pt(10)
            # Alternating row colors
            if row_idx % 2 == 0:
                shading_elm = cell._element.get_or_add_tcPr()
                shading = shading_elm.makeelement(qn("w:shd"), {
                    qn("w:fill"): "D6E4F0",
                    qn("w:val"): "clear",
                })
                shading_elm.append(shading)

    doc.add_paragraph()  # spacing
    return table


def add_rag_status(doc, status, text):
    """Add a RAG (Red/Amber/Green) status indicator."""
    para = doc.add_paragraph()
    colors = {"red": RED, "amber": AMBER, "green": GREEN}
    symbols = {"red": "\u25cf RED", "amber": "\u25cf AMBER", "green": "\u25cf GREEN"}

    run = para.add_run(symbols.get(status.lower(), "\u25cf") + " ")
    run.font.color.rgb = colors.get(status.lower(), ACCENT_GRAY)
    run.font.bold = True
    run.font.size = Pt(11)

    run2 = para.add_run(text)
    run2.font.size = Pt(11)
    run2.font.color.rgb = ACCENT_GRAY


def add_action_item(doc, item_num, description, owner, deadline, status="Open"):
    """Add a formatted action item."""
    para = doc.add_paragraph()
    run = para.add_run(f"AI-{item_num:03d}: ")
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE

    run2 = para.add_run(f"{description}\n")
    run2.font.color.rgb = ACCENT_GRAY

    run3 = para.add_run(f"    Owner: {owner} | Deadline: {deadline} | Status: {status}")
    run3.font.size = Pt(9)
    run3.font.italic = True
    run3.font.color.rgb = ACCENT_GRAY


def generate_status_report(title, project_name, status_data, output_path):
    """Generate a complete project status report."""
    doc = create_base_document()
    add_title_page(doc, title, subtitle=f"Project: {project_name}")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(status_data.get("summary", "Status report summary."))

    doc.add_heading("Overall Status", level=2)
    add_rag_status(doc, status_data.get("rag", "green"), status_data.get("rag_text", "On track"))

    if "milestones" in status_data:
        doc.add_heading("Milestones", level=2)
        add_styled_table(doc, ["Milestone", "Target Date", "Status"], status_data["milestones"])

    if "action_items" in status_data:
        doc.add_heading("Action Items", level=1)
        for i, item in enumerate(status_data["action_items"], 1):
            add_action_item(doc, i, item["desc"], item["owner"], item["deadline"], item.get("status", "Open"))

    if "risks" in status_data:
        doc.add_heading("Risks & Issues", level=1)
        add_styled_table(doc, ["Risk", "Impact", "Likelihood", "Mitigation"], status_data["risks"])

    doc.add_heading("Next Steps", level=1)
    for step in status_data.get("next_steps", ["Define next steps."]):
        doc.add_paragraph(step, style="List Bullet")

    doc.save(output_path)
    print(f"Report generated: {output_path}")
    return output_path


def generate_meeting_minutes(title, meeting_data, output_path):
    """Generate formatted meeting minutes."""
    doc = create_base_document()
    add_title_page(doc, "Meeting Minutes", subtitle=title)

    doc.add_heading("Meeting Details", level=1)
    details = [
        ["Date", meeting_data.get("date", datetime.now().strftime("%B %d, %Y"))],
        ["Time", meeting_data.get("time", "TBD")],
        ["Location", meeting_data.get("location", "TBD")],
        ["Attendees", ", ".join(meeting_data.get("attendees", []))],
    ]
    add_styled_table(doc, ["Field", "Details"], details)

    doc.add_heading("Discussion Points", level=1)
    for point in meeting_data.get("discussion", []):
        doc.add_heading(point["topic"], level=2)
        doc.add_paragraph(point["notes"])

    if "decisions" in meeting_data:
        doc.add_heading("Decisions Made", level=1)
        for dec in meeting_data["decisions"]:
            doc.add_paragraph(dec, style="List Bullet")

    if "action_items" in meeting_data:
        doc.add_heading("Action Items", level=1)
        for i, item in enumerate(meeting_data["action_items"], 1):
            add_action_item(doc, i, item["desc"], item["owner"], item["deadline"])

    doc.save(output_path)
    print(f"Meeting minutes generated: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Strategic Report Generator ready.")
