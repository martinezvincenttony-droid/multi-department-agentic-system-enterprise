"""
Strategic Report Generator — produces professional .docx documents.
Uses python-docx (preinstalled).
"""
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x78)
GREY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _setup_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    return p


def _banner(doc, title, subtitle=None):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade(cell, "1F4E78")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = WHITE
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(11)
        r2.font.color.rgb = WHITE
    doc.add_paragraph()


def _shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _footer(doc, text):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GREY


# ---------- Executive Summary ----------
def executive_summary(title, audience, key_findings, recommendations, next_steps, out_path):
    doc = _setup_doc()
    _banner(doc, title, f"Executive Summary | Prepared for: {audience}")
    _heading(doc, "Key Findings", 2)
    for f in key_findings:
        doc.add_paragraph(f, style="List Bullet")
    _heading(doc, "Recommendations", 2)
    for r in recommendations:
        doc.add_paragraph(r, style="List Bullet")
    _heading(doc, "Next Steps", 2)
    for n in next_steps:
        doc.add_paragraph(n, style="List Number")
    _footer(doc, f"Generated {datetime.now().strftime('%Y-%m-%d')} | Confidential")
    doc.save(out_path)
    return out_path


# ---------- Status Report (RAG) ----------
RAG_COLORS = {"GREEN": "70AD47", "AMBER": "FFC000", "RED": "C00000"}
RAG_EMOJI = {"GREEN": "🟢", "AMBER": "🟡", "RED": "🔴"}


def status_report(project, overall_status, workstreams, blockers, milestones, out_path):
    """
    workstreams: list of {name, status, progress_pct, notes}
    milestones: list of {name, due, status}
    """
    doc = _setup_doc()
    _banner(doc, f"{project} — Status Report", datetime.now().strftime("%B %d, %Y"))

    _heading(doc, f"Overall Status: {RAG_EMOJI.get(overall_status,'')} {overall_status}", 1)

    _heading(doc, "Workstreams", 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Workstream", "Status", "Progress", "Notes"]):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
    for w in workstreams:
        row = table.add_row().cells
        row[0].text = w["name"]
        row[1].text = f"{RAG_EMOJI.get(w['status'],'')} {w['status']}"
        row[2].text = f"{w.get('progress_pct',0)}%"
        row[3].text = w.get("notes", "")

    _heading(doc, "Blockers & Risks", 2)
    if blockers:
        for b in blockers:
            doc.add_paragraph(b, style="List Bullet")
    else:
        doc.add_paragraph("None at this time.")

    _heading(doc, "Upcoming Milestones", 2)
    mt = doc.add_table(rows=1, cols=3)
    mt.style = "Light Grid Accent 1"
    for i, h in enumerate(["Milestone", "Due", "Status"]):
        mt.rows[0].cells[i].text = h
    for m in milestones:
        row = mt.add_row().cells
        row[0].text = m["name"]
        row[1].text = m["due"]
        row[2].text = m["status"]

    _footer(doc, f"{project} Status | {datetime.now().strftime('%Y-%m-%d')}")
    doc.save(out_path)
    return out_path


# ---------- Meeting Minutes ----------
def meeting_minutes(meeting_title, date, attendees, agenda_items, decisions, action_items, out_path):
    """action_items: list of {task, owner, due, status}"""
    doc = _setup_doc()
    _banner(doc, meeting_title, f"{date}  |  Meeting Minutes")
    _heading(doc, "Attendees", 2)
    doc.add_paragraph(", ".join(attendees))

    _heading(doc, "Agenda & Discussion", 2)
    for item in agenda_items:
        p = doc.add_paragraph()
        p.add_run(f"• {item['topic']}: ").bold = True
        p.add_run(item["notes"])

    _heading(doc, "Decisions Made", 2)
    for d in decisions:
        doc.add_paragraph(d, style="List Bullet")

    _heading(doc, "Action Items", 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Task", "Owner", "Due", "Status"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for a in action_items:
        row = table.add_row().cells
        row[0].text = a["task"]
        row[1].text = a["owner"]
        row[2].text = a["due"]
        row[3].text = a.get("status", "Open")

    _footer(doc, f"Minutes recorded {datetime.now().strftime('%Y-%m-%d')}")
    doc.save(out_path)
    return out_path


# ---------- Action Item Tracker ----------
def action_tracker(title, items, out_path):
    """items: list of {task, owner, due, priority, status}"""
    doc = _setup_doc()
    _banner(doc, title, "Action Item Tracker")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    headers = ["Task", "Owner", "Due", "Priority", "Status"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    # Sort by priority then due
    prio_rank = {"P1": 1, "P2": 2, "P3": 3, "P4": 4}
    items_sorted = sorted(items, key=lambda x: (prio_rank.get(x.get("priority", "P3"), 5), x.get("due", "")))
    for it in items_sorted:
        row = table.add_row().cells
        row[0].text = it["task"]
        row[1].text = it["owner"]
        row[2].text = it["due"]
        row[3].text = it.get("priority", "P3")
        row[4].text = it.get("status", "Open")

    _footer(doc, f"Generated {datetime.now().strftime('%Y-%m-%d')} | {len(items)} open items")
    doc.save(out_path)
    return out_path


# ---------- Weekly Digest ----------
def weekly_digest(week_label, accomplishments, in_progress, upcoming, metrics, out_path):
    doc = _setup_doc()
    _banner(doc, "Weekly Digest", week_label)

    _heading(doc, "✅ Accomplishments", 2)
    for a in accomplishments:
        doc.add_paragraph(a, style="List Bullet")

    _heading(doc, "🔄 In Progress", 2)
    for p in in_progress:
        doc.add_paragraph(p, style="List Bullet")

    _heading(doc, "📅 Upcoming", 2)
    for u in upcoming:
        doc.add_paragraph(u, style="List Bullet")

    if metrics:
        _heading(doc, "📊 Key Metrics", 2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Value"
        for k, v in metrics.items():
            row = table.add_row().cells
            row[0].text = k
            row[1].text = str(v)

    _footer(doc, f"Weekly Digest | {week_label}")
    doc.save(out_path)
    return out_path
